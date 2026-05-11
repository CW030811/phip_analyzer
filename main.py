"""PHIP Analyzer CLI

用法：
  python main.py scan
      # 扫描 HKEX 当前 active PHIP 列表，发现新增项目并入库

  python main.py run
      # 处理所有 pending（已发现但未分析）的 PHIP，下载 -> 解析 -> 分析 -> 生成报告

  python main.py analyze --pdf-url <url> --company "公司名" --stock-code "1234"
      # 直接处理一个指定 PHIP（绕过监控，用于测试或补漏）

  python main.py analyze --pdf-path <path> --company "公司名"
      # 处理本地已下载的 PDF

  python main.py status
      # 显示当前数据库中各项目的状态
"""
from __future__ import annotations

import hashlib
import logging
import sys
from datetime import datetime, date
from pathlib import Path

import click
from rich.console import Console
from rich.table import Table
from rich.logging import RichHandler

from src.analyzer import run_analysis
from src.config import get_config
from src.db import PhipDB
from src.downloader import download_pdf
from src.emailer import EmailConfig, send_email
from src.monitor import scan_new_phips
from src.pdf_parser import parse_phip, cache_parsed
from src.report_builder import build_report


console = Console()


def _setup_logging(level: str = "INFO") -> None:
    logging.basicConfig(
        level=getattr(logging, level),
        format="%(message)s",
        datefmt="[%X]",
        handlers=[RichHandler(console=console, rich_tracebacks=True,
                              show_path=False)],
    )


def _short_company(name: str) -> str:
    if not name:
        return "unknown"
    # 取前 8 个字符，去除特殊字符
    import re as _re
    return _re.sub(r'[\\/:*?"<>|\s]', "", name)[:12] or "unknown"


def _filename_for_report(*, cfg, company: str, stock_code: str | None) -> str:
    template = cfg.get("report", "filename_template",
                       default="{date}_{stock_code}_{company_short}_{model_tag}_深度研报.docx")
    # 主模型 tag（去除特殊字符，截短）
    model = str(cfg.get("analyzer", "model", default="model"))
    model_tag = model.replace(":", "-").replace("/", "-")[:24]
    return template.format(
        date=datetime.now().strftime("%Y%m%d"),
        stock_code=stock_code or "NA",
        company_short=_short_company(company),
        model_tag=model_tag,
    )


def _daily_email_body(*, discovered: list[dict], processed: list[dict],
                      failed: list[dict], skipped: int) -> str:
    lines = [
        "PHIP Analyzer 每日自动播报",
        "",
        f"运行日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"新发现 PHIP：{len(discovered)}",
        f"本次生成深度研报：{len(processed)}",
        f"失败：{len(failed)}",
        f"未处理 pending：{skipped}",
        "",
    ]

    if discovered:
        lines.append("一、新发现")
        for item in discovered:
            lines.append(
                f"- {item.get('company_name') or '未知公司'} | "
                f"{item.get('publish_date') or '未知日期'} | {item.get('source_url')}"
            )
        lines.append("")

    if processed:
        lines.append("二、本次生成报告")
        for item in processed:
            lines.append(
                f"- {item.get('company_name') or '未知公司'} | "
                f"{item.get('report_path')}"
            )
        lines.append("")

    if failed:
        lines.append("三、失败项")
        for item in failed:
            lines.append(
                f"- {item.get('company_name') or '未知公司'} | "
                f"{item.get('error')}"
            )
        lines.append("")

    if not discovered and not processed and not failed:
        lines.append("今日没有新增 PHIP，也没有待处理报告。任务已正常运行。")
        lines.append("")

    lines.append("附件为本次新生成的 Word 研报；若无附件，说明今日无需生成新深度报告。")
    return "\n".join(lines)


def _process_one(*, cfg, db: PhipDB, source_url: str | None,
                 pdf_path: str | None, company: str | None,
                 stock_code: str | None, board: str = "main",
                 doc_type: str = "PHIP") -> Path | None:
    """端到端处理一个 PHIP：下载 -> 解析 -> 分析 -> 报告"""
    # 1. 准备数据库记录
    if source_url:
        existing = db.get(source_url)
        if not existing:
            db.upsert_discovered(source_url=source_url, company_name=company,
                                 stock_code=stock_code, board=board,
                                 document_type=doc_type, publish_date=None,
                                 sponsor=None)

    # 2. 获取 PDF
    if pdf_path:
        local_path = Path(pdf_path)
        if not local_path.exists():
            console.print(f"[red]找不到本地 PDF: {pdf_path}[/red]")
            return None
        pdf_info = {"path": str(local_path),
                    "size_bytes": local_path.stat().st_size,
                    "sha256": hashlib.sha256(local_path.read_bytes()).hexdigest()}
        run_key = pdf_info["sha256"][:16]
    elif source_url:
        console.print(f"[cyan]下载 PDF...[/cyan] {source_url}")
        pdf_info = download_pdf(
            url=source_url, stock_code=stock_code, company=company,
            pdf_dir=cfg.get("storage", "pdf_dir"),
            timeout=cfg.get("downloader", "timeout", default=300),
            max_size_mb=cfg.get("downloader", "max_file_size_mb", default=200),
            user_agent=cfg.get("monitor", "user_agent", default="Mozilla/5.0"),
        )
        run_key = pdf_info["sha256"][:16]
        if source_url:
            db.update(source_url, pdf_path=pdf_info["path"],
                      pdf_size_bytes=pdf_info["size_bytes"],
                      status="DOWNLOADED")
    else:
        console.print("[red]必须提供 --pdf-url 或 --pdf-path[/red]")
        return None

    # 3. 解析
    console.print("[cyan]解析 PDF 结构...[/cyan]")
    parsed = parse_phip(pdf_info["path"], cfg)
    cache_parsed(parsed, cfg.get("storage", "cache_dir"), run_key)
    if source_url:
        db.update(source_url, pdf_pages=parsed.total_pages, status="PARSED")
    console.print(f"  → {parsed.total_pages} 页, "
                  f"{len(parsed.sections)} 个标准章节命中")

    # 4. 分析
    console.print("[cyan]启动 Claude API 分析（这一步约 3-8 分钟，按 PDF 大小）...[/cyan]")
    cache_dir = Path(cfg.get("storage", "cache_dir")) / "analysis"
    result = run_analysis(parsed, config=cfg, cache_dir=cache_dir,
                          run_key=run_key)
    failed_stages = {
        name: status
        for name, status in (result.section_status or {}).items()
        if str(status).startswith("failed")
    }
    if failed_stages:
        raise RuntimeError(f"analysis stages failed: {failed_stages}")

    # 招股书前几页通常是「警告／重要提示」，公司名常需用户/数据库提供作为兜底
    if (not result.company_name or result.company_name == "未知公司") and company:
        result.company_name = company
        result.cover_info.setdefault("company_name_cn", company)

    if source_url:
        db.update(source_url, status="ANALYZED",
                  company_name=result.company_name)

    # 5. 报告
    fname = _filename_for_report(cfg=cfg, company=result.company_name,
                                 stock_code=stock_code or
                                 result.cover_info.get("stock_code"))
    out_path = Path(cfg.get("report", "output_dir")) / fname
    console.print(f"[cyan]生成 Word 报告...[/cyan]")
    build_report(result, parsed.to_dict(), output_path=out_path)

    if source_url:
        db.update(source_url, status="REPORTED",
                  report_path=str(out_path))
    console.print(f"[green]✓ 完成: {out_path}[/green]")
    return out_path


# ============ CLI 命令 ============

@click.group()
@click.option("--log-level", default=None, help="覆盖日志级别 (DEBUG/INFO/WARNING)")
@click.pass_context
def cli(ctx, log_level: str | None):
    cfg = get_config()
    _setup_logging(log_level or cfg.log_level)
    ctx.ensure_object(dict)
    ctx.obj["cfg"] = cfg
    ctx.obj["db"] = PhipDB(cfg.get("storage", "db_path"))


@cli.command()
@click.pass_context
def scan(ctx):
    """扫描 HKEX，发现新增 PHIP 并入库（不下载、不分析）"""
    cfg = ctx.obj["cfg"]
    db = ctx.obj["db"]
    run_id = db.start_run()
    try:
        new_items = scan_new_phips(cfg, db)
        db.finish_run(run_id, discovered=len(new_items),
                      analyzed=0, failed=0,
                      notes=f"scan only, found {len(new_items)} new")
        console.print(f"[green]扫描完成。新增 {len(new_items)} 条。[/green]")
        for it in new_items:
            console.print(f"  - {it.get('company_name')}  ({it['source_url']})")
    except Exception as e:
        db.finish_run(run_id, discovered=0, analyzed=0, failed=1,
                      notes=f"error: {e}")
        raise


@cli.command()
@click.option("--max-items", default=10, type=int,
              help="本轮最多处理多少个 pending 项目")
@click.option("--since-days", default=None, type=int,
              help="只处理 publish_date 在过去 N 天内的项目（如 7 表示本周）")
@click.pass_context
def run(ctx, max_items: int, since_days: int | None):
    """完整流程：扫描 -> 处理所有 pending 项目"""
    cfg = ctx.obj["cfg"]
    db = ctx.obj["db"]
    run_id = db.start_run()
    discovered = 0
    analyzed = 0
    failed = 0

    try:
        # 1. 扫描
        new_items = scan_new_phips(cfg, db)
        discovered = len(new_items)

        # 2. 处理 pending（按 publish_date 倒序，可选 since_days 过滤）
        pending = db.list_pending(
            statuses=("DISCOVERED", "DOWNLOADED", "PARSED"),
            since_days=since_days,
        )
        if since_days:
            console.print(f"[cyan]待处理项目（最近 {since_days} 天内）: "
                          f"{len(pending)}[/cyan]")
        else:
            console.print(f"[cyan]待处理项目: {len(pending)}[/cyan]")

        for record in pending[:max_items]:
            console.rule(f"[bold cyan]{record.get('company_name')} "
                         f"({record.get('publish_date') or '?'})")
            try:
                _process_one(
                    cfg=cfg, db=db,
                    source_url=record["source_url"],
                    pdf_path=record.get("pdf_path"),
                    company=record.get("company_name"),
                    stock_code=record.get("stock_code"),
                    board=record.get("board") or "main",
                    doc_type=record.get("document_type") or "PHIP",
                )
                analyzed += 1
            except Exception as e:
                logging.exception("处理失败: %s", record["source_url"])
                db.mark_failed(record["source_url"], str(e))
                failed += 1
                console.print(f"[red]✗ 失败: {e}[/red]")
    finally:
        db.finish_run(run_id, discovered=discovered, analyzed=analyzed,
                      failed=failed)
        console.print(f"\n[bold]本轮汇总：发现 {discovered} 新 / "
                      f"成功 {analyzed} / 失败 {failed}[/bold]")


@cli.command()
@click.option("--pdf-url", help="HKEX PHIP 的 PDF URL（远程下载）")
@click.option("--pdf-path", help="本地 PDF 路径（直接处理）")
@click.option("--company", help="公司名称")
@click.option("--stock-code", help="拟上市代码")
@click.option("--board", default="main", type=click.Choice(["main", "gem"]),
              help="主板 / 创业板")
@click.pass_context
def analyze(ctx, pdf_url: str, pdf_path: str, company: str,
            stock_code: str, board: str):
    """处理单个指定 PHIP（用于测试或补漏）"""
    cfg = ctx.obj["cfg"]
    db = ctx.obj["db"]
    if not pdf_url and not pdf_path:
        console.print("[red]必须提供 --pdf-url 或 --pdf-path[/red]")
        sys.exit(1)
    _process_one(
        cfg=cfg, db=db,
        source_url=pdf_url, pdf_path=pdf_path,
        company=company, stock_code=stock_code, board=board,
        doc_type="PHIP",
    )


@cli.command()
@click.pass_context
def status(ctx):
    """显示数据库中的项目状态"""
    db = ctx.obj["db"]
    with db.conn() as c:
        rows = c.execute(
            """SELECT status, COUNT(*) as n FROM phip
               GROUP BY status ORDER BY n DESC"""
        ).fetchall()
        recent = c.execute(
            """SELECT company_name, stock_code, document_type, status,
                      report_path, updated_at
               FROM phip ORDER BY updated_at DESC LIMIT 20"""
        ).fetchall()

    t1 = Table(title="状态汇总", show_lines=False)
    t1.add_column("状态")
    t1.add_column("数量", justify="right")
    for r in rows:
        t1.add_row(r["status"], str(r["n"]))
    console.print(t1)

    t2 = Table(title="近 20 条记录", show_lines=False)
    for col in ["公司", "代码", "文档", "状态", "报告", "更新时间"]:
        t2.add_column(col)
    for r in recent:
        t2.add_row(
            (r["company_name"] or "")[:25],
            r["stock_code"] or "—",
            r["document_type"] or "",
            r["status"] or "",
            "✓" if r["report_path"] else "",
            r["updated_at"] or "",
        )
    console.print(t2)


@cli.command()
@click.option("--scan-first/--no-scan", default=True,
              help="生成前先扫描 HKEX 同步最新 active 列表")
@click.option("--output", default=None, help="输出 Word 文件路径")
@click.pass_context
def report(ctx, scan_first: bool, output: str | None):
    """生成"当前 active PHIP 列表"汇报 Word 文档（不调用 LLM）"""
    from datetime import datetime
    from urllib.parse import urlparse
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    cfg = ctx.obj["cfg"]
    db = ctx.obj["db"]

    if scan_first:
        console.print("[cyan]同步 HKEX active PHIP 列表...[/cyan]")
        scan_new_phips(cfg, db)

    with db.conn() as c:
        rows = c.execute(
            """SELECT company_name, stock_code, board, document_type,
                      publish_date, source_url, status, report_path,
                      pdf_pages, discovered_at
               FROM phip
               WHERE status != 'FAILED'
               ORDER BY COALESCE(publish_date, substr(discovered_at,1,10)) DESC"""
        ).fetchall()

    if not rows:
        console.print("[yellow]数据库为空，先运行 python main.py scan[/yellow]")
        return

    today = datetime.now().strftime("%Y%m%d")
    out_path = Path(output) if output else (
        Path(cfg.get("report", "output_dir")) / f"{today}_HKEX_Active_PHIP_汇报.docx"
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("港交所 Active PHIP 项目汇报")
    r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}    "
                  f"项目数：{len(rows)}    数据源：HKEXnews appactive")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    headers = ["#", "公司", "板块", "文档", "递交日期", "状态", "已生成研报"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = table.cell(0, j); cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.size = Pt(10); run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)

    for i, row in enumerate(rows, 1):
        cells = [
            str(i),
            row["company_name"] or "—",
            "主板" if row["board"] == "main" else "创业板",
            row["document_type"] or "—",
            row["publish_date"] or "—",
            row["status"] or "—",
            "✓" if row["report_path"] else "—",
        ]
        for j, val in enumerate(cells):
            c2 = table.cell(i, j); c2.text = ""
            run = c2.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)

    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run("项目源链接")
    r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    for i, row in enumerate(rows, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {row['company_name'] or '—'}")
        run.font.size = Pt(10); run.bold = True
        p2 = doc.add_paragraph()
        run = p2.add_run(row["source_url"] or "")
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = "Consolas"
        if row["report_path"]:
            p3 = doc.add_paragraph()
            run = p3.add_run(f"   研报：{row['report_path']}")
            run.font.size = Pt(9); run.italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("说明：本汇报基于 HKEX 公开 active 列表实时拉取并经去重（同公司中英版本仅保留中文）。"
                  "各公司单独深度研报通过 `python main.py run` 生成。")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(str(out_path))
    console.print(f"[green]✓ 汇报已生成：{out_path}[/green]")

@cli.command("test-email")
@click.option("--to", default=None, help="Override recipient for this test")
@click.option("--dry-run", is_flag=True, help="Validate config without sending")
def test_email(to: str | None, dry_run: bool):
    """Validate SMTP settings and optionally send a small test email."""
    cfg = EmailConfig.from_env(default_to=to)
    if to:
        cfg = EmailConfig(
            host=cfg.host,
            port=cfg.port,
            use_ssl=cfg.use_ssl,
            username=cfg.username,
            password=cfg.password,
            sender=cfg.sender,
            sender_name=cfg.sender_name,
            recipients=[to],
        )

    if dry_run:
        console.print(
            f"[green]SMTP config OK[/green] {cfg.username} -> "
            f"{', '.join(cfg.recipients)} via {cfg.host}:{cfg.port}"
        )
        return

    send_email(
        subject=f"PHIP Analyzer 邮件测试 {date.today().isoformat()}",
        body="这是一封 PHIP Analyzer 自动播报测试邮件。收到这封邮件表示 QQ SMTP 配置可用。",
        config=cfg,
    )
    console.print("[green]测试邮件已发送[/green]")


@cli.command()
@click.option("--max-items", default=1, type=int,
              help="Maximum pending PHIP records to analyze in one daily run")
@click.option("--since-days", default=30, type=int,
              help="Only process records published/discovered in recent N days")
@click.option("--send-empty/--no-send-empty", default=True,
              help="Send a heartbeat email even when nothing changed")
@click.option("--to", default=None, help="Override recipient for this run")
@click.pass_context
def daily(ctx, max_items: int, since_days: int, send_empty: bool, to: str | None):
    """Scan, analyze pending PHIPs, and email the daily report."""
    cfg = ctx.obj["cfg"]
    db = ctx.obj["db"]
    with db.conn() as c:
        c.execute(
            """UPDATE run_log
               SET finished_at = ?, failed_count = failed_count + 1,
                   notes = COALESCE(notes || '; ', '') || 'closed stale run'
               WHERE finished_at IS NULL""",
            (datetime.utcnow().isoformat(timespec="seconds"),),
        )
    run_id = db.start_run()
    discovered: list[dict] = []
    processed: list[dict] = []
    failed: list[dict] = []

    try:
        discovered = scan_new_phips(cfg, db)
        pending = db.list_pending(
            statuses=("DISCOVERED", "DOWNLOADED", "PARSED"),
            since_days=since_days,
        )
        to_process = pending[:max_items]
        skipped = max(0, len(pending) - len(to_process))

        for record in to_process:
            console.rule(f"[bold cyan]{record.get('company_name')} "
                         f"({record.get('publish_date') or '?'})")
            try:
                out_path = _process_one(
                    cfg=cfg, db=db,
                    source_url=record["source_url"],
                    pdf_path=record.get("pdf_path"),
                    company=record.get("company_name"),
                    stock_code=record.get("stock_code"),
                    board=record.get("board") or "main",
                    doc_type=record.get("document_type") or "PHIP",
                )
                if out_path:
                    processed.append({
                        "company_name": record.get("company_name"),
                        "source_url": record.get("source_url"),
                        "report_path": str(out_path),
                    })
            except Exception as e:
                logging.exception("daily failed: %s", record.get("source_url"))
                if record.get("source_url"):
                    db.mark_failed(record["source_url"], str(e))
                failed.append({
                    "company_name": record.get("company_name"),
                    "source_url": record.get("source_url"),
                    "error": str(e),
                })

        db.finish_run(run_id, discovered=len(discovered),
                      analyzed=len(processed), failed=len(failed),
                      notes="daily email flow")

        if not (send_empty or discovered or processed or failed):
            console.print("[yellow]No updates and heartbeat email disabled[/yellow]")
            return

        subject = (
            f"PHIP 自动播报 {date.today().isoformat()} | "
            f"新增{len(discovered)} 研报{len(processed)} 失败{len(failed)}"
        )
        body = _daily_email_body(
            discovered=discovered,
            processed=processed,
            failed=failed,
            skipped=skipped,
        )
        attachments = [item["report_path"] for item in processed]
        email_cfg = EmailConfig.from_env(default_to=to)
        if to:
            email_cfg = EmailConfig(
                host=email_cfg.host,
                port=email_cfg.port,
                use_ssl=email_cfg.use_ssl,
                username=email_cfg.username,
                password=email_cfg.password,
                sender=email_cfg.sender,
                sender_name=email_cfg.sender_name,
                recipients=[to],
            )
        send_email(
            subject=subject,
            body=body,
            attachments=attachments,
            config=email_cfg,
        )
        console.print(f"[green]每日播报邮件已发送：{', '.join(email_cfg.recipients)}[/green]")
    except Exception as e:
        db.finish_run(run_id, discovered=len(discovered),
                      analyzed=len(processed), failed=len(failed) + 1,
                      notes=f"daily error: {e}")
        raise


if __name__ == "__main__":
    cli(obj={})
