"""Word 报告生成器：把分析结果组装成 20+ 页深度研报

设计：
- 用 python-docx 生成 .docx
- 内置一个轻量级 Markdown -> docx 渲染器（支持 # 标题、**粗体**、*斜体*、表格、列表）
- 封面 + 目录 + 各章节
"""
from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches

from .analyzer import AnalysisResult

logger = logging.getLogger(__name__)


# ============ 样式常量 ============

FONT_HAN = "微软雅黑"
FONT_LATIN = "Calibri"
COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # 深蓝
COLOR_ACCENT = RGBColor(0xC0, 0x39, 0x2B)    # 红色（用于关注点）
COLOR_GREY = RGBColor(0x55, 0x55, 0x55)


# ============ 辅助：设置中英文字体 ============

def _set_run_font(run, han: str = FONT_HAN, latin: str = FONT_LATIN,
                  size_pt: float | None = None,
                  color: RGBColor | None = None,
                  bold: bool | None = None) -> None:
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), han)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def _add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()


def _add_horizontal_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============ 封面 ============

def _add_cover(doc: Document, *, company_name: str, stock_code: str | None,
               cover_info: dict, run_date: str) -> None:
    # 留白
    for _ in range(4):
        doc.add_paragraph()

    # 报告类型
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("港股 IPO 深度研究")
    _set_run_font(r, size_pt=14, color=COLOR_GREY)

    doc.add_paragraph()

    # 公司名
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(company_name or "（公司名称待确认）")
    _set_run_font(r, size_pt=28, color=COLOR_PRIMARY, bold=True)

    if cover_info.get("company_name_en"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover_info["company_name_en"])
        _set_run_font(r, size_pt=14, color=COLOR_GREY, latin="Calibri")

    doc.add_paragraph()
    doc.add_paragraph()

    # 元信息表
    meta_rows = [
        ("拟上市板块", cover_info.get("listing_board") or "—"),
        ("拟上市代码", stock_code or cover_info.get("stock_code") or "待定"),
        ("一级行业", cover_info.get("industry_primary") or "—"),
        ("二级行业", cover_info.get("industry_secondary") or "—"),
        ("保荐人", ", ".join(cover_info.get("sponsor", []) or []) or "—"),
        ("特殊机制", ", ".join(filter(None, [
            "18A 生物科技" if cover_info.get("is_biotech_18a") else None,
            "特专科技章节" if cover_info.get("is_specialist_tech") else None,
            "WVR 同股不同权" if cover_info.get("is_weighted_voting") else None,
        ])) or "—"),
        ("业绩记录期", cover_info.get("track_record_period") or "—"),
        ("研报日期", run_date),
    ]
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (k, v) in enumerate(meta_rows):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        c1.width = Cm(4.5)
        c2.width = Cm(9)
        c1.text = ""
        c2.text = ""
        r1 = c1.paragraphs[0].add_run(k)
        _set_run_font(r1, size_pt=10.5, color=COLOR_GREY)
        r2 = c2.paragraphs[0].add_run(str(v))
        _set_run_font(r2, size_pt=10.5, bold=True)

    # 免责
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本报告基于港交所公开披露的招股说明书生成，仅供研究参考，不构成投资建议。")
    _set_run_font(r, size_pt=9, color=COLOR_GREY)

    _add_page_break(doc)


# ============ Markdown 渲染器 ============

# 简化的 Markdown 解析器：足够处理 LLM 生成的常规 markdown
TABLE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\-]+[\s|:\-]*$")


def _split_markdown_blocks(md: str) -> Iterator[tuple[str, list[str]]]:
    """把 markdown 切分为 (block_type, lines) 序列"""
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 表格：连续两行以上 |...|，且第二行是分隔行
        if (TABLE_LINE_RE.match(line) and i + 1 < n and
                TABLE_SEP_RE.match(lines[i + 1])):
            tbl: list[str] = []
            while i < n and TABLE_LINE_RE.match(lines[i]):
                tbl.append(lines[i])
                i += 1
            yield "table", tbl
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            yield f"h{len(m.group(1))}", [m.group(2)]
            i += 1
            continue

        # 列表（无序）
        if re.match(r"^[\-\*]\s+", stripped):
            buf = []
            while i < n and re.match(r"^[\-\*]\s+", lines[i].strip()):
                buf.append(re.sub(r"^[\-\*]\s+", "", lines[i].strip()))
                i += 1
            yield "ul", buf
            continue

        # 列表（有序）
        if re.match(r"^\d+\.\s+", stripped):
            buf = []
            while i < n and re.match(r"^\d+\.\s+", lines[i].strip()):
                buf.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            yield "ol", buf
            continue

        # 段落：连续非空、非列表/标题/表格行
        buf = []
        while i < n and lines[i].strip() and not (
                re.match(r"^#{1,6}\s+", lines[i].strip()) or
                re.match(r"^[\-\*]\s+", lines[i].strip()) or
                re.match(r"^\d+\.\s+", lines[i].strip()) or
                TABLE_LINE_RE.match(lines[i])):
            buf.append(lines[i])
            i += 1
        yield "p", buf


def _render_inline(p, text: str) -> None:
    """处理 **粗体** *斜体*"""
    pos = 0
    pattern = re.compile(r"(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            _set_run_font(r, size_pt=10.5)
        token = m.group(0)
        if token.startswith("**"):
            r = p.add_run(token[2:-2])
            _set_run_font(r, size_pt=10.5, bold=True)
        elif token.startswith("`"):
            r = p.add_run(token[1:-1])
            _set_run_font(r, size_pt=10.5, latin="Consolas")
            r.font.color.rgb = COLOR_GREY
        else:
            r = p.add_run(token[1:-1])
            _set_run_font(r, size_pt=10.5)
            r.italic = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        _set_run_font(r, size_pt=10.5)


def _add_heading(doc: Document, text: str, level: int) -> None:
    sizes = {1: 18, 2: 15, 3: 13, 4: 11.5, 5: 11, 6: 10.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _set_run_font(r, size_pt=sizes.get(level, 11), color=COLOR_PRIMARY, bold=True)


def _add_paragraph(doc: Document, lines: list[str]) -> None:
    text = " ".join(line.strip() for line in lines if line.strip())
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    _render_inline(p, text)


def _add_list(doc: Document, items: list[str], ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        _render_inline(p, it)


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for ln in lines:
        if TABLE_SEP_RE.match(ln):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table(doc: Document, lines: list[str]) -> None:
    rows = _parse_md_table(lines)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # 补齐
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            _render_inline(p, cell_text)
            for r in p.runs:
                _set_run_font(r, size_pt=9.5,
                              bold=(i == 0),
                              color=(RGBColor(0xFF, 0xFF, 0xFF) if i == 0 else None))
            if i == 0:
                # 表头底色
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F3A5F")
                tcPr.append(shd)


def render_markdown(doc: Document, md: str) -> None:
    if not md or not md.strip():
        return
    for block_type, lines in _split_markdown_blocks(md):
        if block_type.startswith("h"):
            level = int(block_type[1:])
            _add_heading(doc, lines[0], level)
        elif block_type == "p":
            _add_paragraph(doc, lines)
        elif block_type == "ul":
            _add_list(doc, lines, ordered=False)
        elif block_type == "ol":
            _add_list(doc, lines, ordered=True)
        elif block_type == "table":
            _add_table(doc, lines)


# ============ 报告组装 ============

CHAPTER_TITLES = {
    "diagonal": "投资建议（横纵交汇）",
    "diachronic": "一、纵向分析：公司发展史",
    "industry": "二、横向分析｜行业速描",
    "business_financial": "三、商业模式与财务穿透",
    "peer": "四、横向分析｜同业深度对比",
    "uop": "五、募集资金用途",
    "risk": "六、风险因素",
    "appendix": "附录：研究方法与免责声明",
}


def _add_chapter_header(doc: Document, title: str) -> None:
    _add_page_break(doc)
    p = doc.add_paragraph()
    r = p.add_run(title)
    _set_run_font(r, size_pt=20, color=COLOR_PRIMARY, bold=True)
    _add_horizontal_line(doc)
    doc.add_paragraph()


def _format_risks_chapter(doc: Document, risks_json: dict) -> None:
    risks = risks_json.get("risks", [])
    if not risks:
        p = doc.add_paragraph("（风险章节抽取失败或为空）")
        return

    # 表格头
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    headers = ["排名", "风险类别", "标题", "可能性", "影响"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        _set_run_font(r, size_pt=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)

    for r in risks:
        row = table.add_row()
        cells = [
            str(r.get("rank", "")),
            r.get("category", ""),
            r.get("title", ""),
            r.get("likelihood", ""),
            r.get("impact", ""),
        ]
        for j, val in enumerate(cells):
            c = row.cells[j]
            c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            _set_run_font(run, size_pt=9.5)

    doc.add_paragraph()
    # 详细描述
    h = doc.add_paragraph()
    r = h.add_run("风险详述")
    _set_run_font(r, size_pt=13, bold=True, color=COLOR_PRIMARY)

    for r in risks:
        p = doc.add_paragraph()
        run = p.add_run(f"#{r.get('rank', '')}  {r.get('title', '')}")
        _set_run_font(run, size_pt=11.5, bold=True, color=COLOR_ACCENT)

        for label, key in [("描述", "description"),
                           ("证据", "evidence"),
                           ("跟踪指标", "monitoring")]:
            val = r.get(key)
            if val:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}：")
                _set_run_font(run, size_pt=10, bold=True, color=COLOR_GREY)
                run = p.add_run(str(val))
                _set_run_font(run, size_pt=10)
        doc.add_paragraph()


# ============ 主函数 ============

def build_report(result: AnalysisResult, parsed_meta: dict, *,
                 output_path: Path) -> Path:
    """组装并写出 Word 报告"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 全局默认字体设置
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HAN)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    cover_info = result.cover_info or {}
    today = date.today().isoformat()

    # --- 封面 ---
    _add_cover(
        doc,
        company_name=result.company_name,
        stock_code=cover_info.get("stock_code"),
        cover_info=cover_info,
        run_date=today,
    )

    # --- 投资建议先行（横纵交汇 - 报告精华段，前置以便快速决策）---
    _add_chapter_header(doc, CHAPTER_TITLES["diagonal"])
    render_markdown(doc, getattr(result, "diagonal_md", "") or "_投资建议未生成_")

    # --- 一、纵向分析：发展史叙事 ---
    _add_chapter_header(doc, CHAPTER_TITLES["diachronic"])
    render_markdown(doc, getattr(result, "diachronic_md", "") or "_纵向分析未生成_")

    # --- 二、行业速描 ---
    _add_chapter_header(doc, CHAPTER_TITLES["industry"])
    render_markdown(doc, result.industry_md or "_行业分析未生成_")

    # --- 三、商业模式与财务穿透 ---
    _add_chapter_header(doc, CHAPTER_TITLES["business_financial"])
    render_markdown(doc, getattr(result, "business_financial_md", "") or
                    "_商业与财务分析未生成_")

    # --- 四、同业深度对比 ---
    _add_chapter_header(doc, CHAPTER_TITLES["peer"])
    # 先输出同业场景判断
    peers_json = getattr(result, "peers_json", {}) or {}
    scenario = peers_json.get("scenario")
    if scenario:
        rationale = peers_json.get("scenario_rationale", "")
        scenario_md = f"**同业场景判定**：场景 {scenario}\n\n{rationale}\n\n"
        render_markdown(doc, scenario_md)
    render_markdown(doc, result.peer_comparison_md or "_同业对比未生成_")

    # --- 五、募投 ---
    _add_chapter_header(doc, CHAPTER_TITLES["uop"])
    render_markdown(doc, result.use_of_proceeds_md or "_募投分析未生成_")

    # --- 六、风险 ---
    _add_chapter_header(doc, CHAPTER_TITLES["risk"])
    _format_risks_chapter(doc, result.risks_json)

    # --- 附录 ---
    _add_chapter_header(doc, CHAPTER_TITLES["appendix"])
    prompt_version = getattr(result, "prompt_version", "—")
    appendix_md = f"""### 研究框架
本报告采用「横纵分析法」（Prompt 版本 {prompt_version}）：
- **纵向分析（Diachronic）**：沿时间轴还原公司发展史，强调因果链与决策逻辑
- **横向分析（Synchronic）**：行业速描 + 商业模式财务 + 同业深度对比（按场景 A/B/C 展开）
- **横纵交汇**：基于纵横信息的投资建议（ECM 分析师视角，直给买入/打新判断）

### 数据来源
- 港交所披露易（HKEXnews）公开披露的 PHIP / Application Proof
- 招股书原始 PDF 文件：{parsed_meta.get('pdf_path', 'N/A')}
- 原始 PDF 共 {parsed_meta.get('total_pages', 'N/A')} 页

### 研究方法
- 用 PyMuPDF 抽取招股书结构化文本
- 调用 Anthropic Claude API（主分析模型: claude-opus-4-7）做多阶段深度分析
- 五个独立阶段（纵向/行业/商业财务/风险/募投）并发执行，
  之后串行执行同业场景判断 → 同业对比 → 横纵交汇

### 局限性
- 同业财务数据基于模型对公开市场的认知，存在时效性误差，关键数字以 Wind / Bloomberg / 同业最新财报为准
- LLM 生成内容可能存在幻觉，重要决策前请回到招股书原文核对
- 招股书"行业概览"章节为公司方付费报告，倾向乐观，本报告已尝试做批判性重构
- 投资建议为模型基于公开信息的研究输出，**不构成投资建议**

### 章节生成状态
"""
    for k, v in (result.section_status or {}).items():
        appendix_md += f"- {k}: {v}\n"

    appendix_md += """

### 免责声明
本报告由 PHIP Analyzer 自动化工具基于港交所公开披露的 PHIP 文件生成，仅作研究参考，
不构成证券投资建议或要约邀请。投资者应基于独立判断作出决策，并承担相应风险。
"""
    render_markdown(doc, appendix_md)

    doc.save(str(output_path))
    logger.info("报告已生成: %s", output_path)
    return output_path
